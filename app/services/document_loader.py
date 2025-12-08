import logging
import os
import tempfile

from docx import Document
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader

logger = logging.getLogger("app.document_loader")


def _persist_temp(file: UploadFile) -> str:
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as tmp:
            tmp.write(file.file.read())
            return tmp.name
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to store upload: {exc}")


def _cleanup(path: str) -> None:
    try:
        os.remove(path)
    except OSError:
        pass


def read_text(file: UploadFile) -> str:
    suffix = (file.filename or "").lower()
    path = _persist_temp(file)
    try:
        if suffix.endswith(".pdf"):
            reader = PdfReader(path)
            text = "\n".join(filter(None, (page.extract_text() for page in reader.pages)))
            logger.info("Parsed PDF: name=%s pages=%s chars=%s", file.filename, len(reader.pages), len(text))
            return text
        if suffix.endswith(".docx"):
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            logger.info("Parsed DOCX: name=%s paragraphs=%s chars=%s", file.filename, len(doc.paragraphs), len(text))
            return text
        raise HTTPException(status_code=400, detail="Unsupported file type; use .pdf or .docx")
    finally:
        _cleanup(path)
